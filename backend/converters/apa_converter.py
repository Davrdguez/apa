from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def convert_to_apa(input_path, output_path, titulo, autor, institucion, carrera, profesor, ubicacion, fecha, referencias):
    original = Document(input_path)
    new_doc = Document()

    # Márgenes APA
    for section in new_doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Portada APA
    portada_info = [
        {"texto": titulo, "bold": True},
        {"texto": autor},
        {"texto": institucion},
        {"texto": carrera},
        {"texto": profesor},
        {"texto": ubicacion},
        {"texto": fecha}
    ]

    for bloque in portada_info:
        p = new_doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)

        if "texto" in bloque:
            run = p.add_run(str(bloque["texto"]))
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(12)
            run.bold = bloque.get("bold", False)

    new_doc.add_page_break()

    # Detectar títulos
    def es_titulo(paragraph):
        if len(paragraph.runs) == 1:
            run = paragraph.runs[0]
            text = run.text.strip()
            if len(text.split()) <= 10 and (run.bold or run.underline or paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER):
                return True
        return False

    # Contenido principal
    for paragraph in original.paragraphs:
        text = paragraph.text.strip()
        if not text:
            new_doc.add_paragraph()
            continue

        new_p = new_doc.add_paragraph()
        new_p.paragraph_format.line_spacing = 2.0

        if es_titulo(paragraph):
            new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p.paragraph_format.first_line_indent = Inches(0)
            run = new_p.add_run(text.capitalize())
            run.bold = True
        else:
            new_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            new_p.paragraph_format.first_line_indent = Inches(0.5)
            run = new_p.add_run(text)

        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.underline = False
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        if run.font.color:
            run.font.color.rgb = None

    # ---- REFERENCIAS APA ---- #

    def formatear_referencia_libre(texto):
        palabras = texto.strip().split()
        if len(palabras) < 5:
            return texto

        try:
            nombre = palabras[0]
            apellido = palabras[1]
            año = palabras[2]

            # Detectar palabra "editorial"
            idx = -1
            for i in range(3, len(palabras) - 1):
                if palabras[i].lower() == "editorial":
                    idx = i
                    break

            if idx != -1:
                titulo = " ".join(palabras[3:idx])
                editorial = " ".join(palabras[idx:idx+2])
            else:
                titulo = " ".join(palabras[3:-1])
                editorial = palabras[-1]

            autor_formateado = f"{apellido.capitalize()}, {nombre[0].upper()}."
            titulo_formateado = titulo.capitalize()
            editorial_formateado = editorial[0].upper() + editorial[1:]

            return f"{autor_formateado} ({año}). *{titulo_formateado}*. {editorial_formateado}."
        except Exception:
            return texto



    referencias_lista = [
        ref.strip() for ref in referencias.splitlines()
        if ref.strip()
    ]

    if referencias_lista:
        # Título "Referencias"
        ref_title = new_doc.add_paragraph()
        ref_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = ref_title.add_run("Referencias")
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        ref_title.paragraph_format.line_spacing = 2.0

        referencias_formateadas = []
        for ref in referencias_lista:
            formato = formatear_referencia_libre(ref)
            referencias_formateadas.append(formato)

        referencias_formateadas.sort()

        for ref in referencias_formateadas:
            p = new_doc.add_paragraph(ref)
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            run = p.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    new_doc.save(output_path)
